"""
LotusAI Memory Management System
SÃ¼rÃ¼m: 2.5.4
AÃ§Ä±klama: Hybrid hafÄ±za sistemi (SQLite + ChromaDB) ve RAG implementasyonu

Ã–zellikler:
- KÄ±sa sÃ¼reli hafÄ±za: SQLite (hÄ±zlÄ± eriÅŸim)
- Uzun sÃ¼reli hafÄ±za: ChromaDB (anlamsal arama)
- RAG: Belge indeksleme ve retrieval
- GPU hÄ±zlandÄ±rma
- Thread-safe operasyonlar
- Ollama embedding desteÄŸi (nomic-embed-text) â€” AI_PROVIDER='ollama' iÃ§in
"""

import sqlite3
import logging
import threading
import hashlib
import re
import requests
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Any, Optional
from dataclasses import dataclass
from contextlib import contextmanager
from enum import Enum

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# CONFIG
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
from config import Config

logger = logging.getLogger("LotusAI.Memory")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# EXTERNAL LIBRARIES
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# PDF
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("âš ï¸ PyPDF2 yÃ¼klÃ¼ deÄŸil, PDF desteÄŸi yok")

# DOCX
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("âš ï¸ python-docx yÃ¼klÃ¼ deÄŸil, DOCX desteÄŸi yok")

# ChromaDB + Torch
try:
    import chromadb
    from chromadb.config import Settings
    from chromadb.utils import embedding_functions
    import torch
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    logger.warning("âš ï¸ ChromaDB/torch yÃ¼klÃ¼ deÄŸil, vektÃ¶r arama devre dÄ±ÅŸÄ±")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# OLLAMA EMBEDDING FUNCTION
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class OllamaEmbeddingFunction:
    """
    ChromaDB uyumlu Ollama embedding fonksiyonu.

    Config.LOCAL_VEK modelini (varsayÄ±lan: nomic-embed-text) kullanarak
    Ollama'nÄ±n /api/embeddings endpoint'inden vektÃ¶r Ã¼retir.

    KullanÄ±m: AI_PROVIDER='ollama' olduÄŸunda SentenceTransformer'Ä±n yerini alÄ±r.
    """

    def __init__(self, model: str, ollama_url: str, timeout: int = 30):
        """
        Args:
            model: Ollama embedding model adÄ± (Ã¶rn. nomic-embed-text)
            ollama_url: Ollama base URL (Ã¶rn. http://localhost:11434/api)
            timeout: Ä°stek zaman aÅŸÄ±mÄ± (saniye)
        """
        self.model = model
        self.timeout = timeout

        # /api/embeddings endpoint'ini doÄŸru ÅŸekilde oluÅŸtur
        base = ollama_url.rstrip("/")
        if base.endswith("/api"):
            self.endpoint = f"{base}/embeddings"
        elif "/api" in base:
            self.endpoint = base.rsplit("/", 1)[0] + "/embeddings"
        else:
            self.endpoint = f"{base}/api/embeddings"

        logger.info(
            f"ğŸ”¢ OllamaEmbeddingFunction hazÄ±r | "
            f"Model: {self.model} | Endpoint: {self.endpoint}"
        )

    def name(self) -> str:
        """
        ChromaDB'nin gÃ¼ncel sÃ¼rÃ¼mleri, Ã¶zel (custom) embedding fonksiyonlarÄ±nÄ±n
        kimliÄŸini tespit etmek iÃ§in Ã§aÄŸrÄ±labilir bir 'name()' metodu bekler.
        """
        return f"OllamaEmbeddingFunction-{self.model}"

    def __call__(self, input: List[str]) -> List[List[float]]:
        """
        ChromaDB'nin beklediÄŸi arayÃ¼z.

        Args:
            input: VektÃ¶rlenecek metin listesi

        Returns:
            Embedding vektÃ¶rleri listesi
        """
        embeddings = []

        for text in input:
            try:
                response = requests.post(
                    self.endpoint,
                    json={"model": self.model, "prompt": text},
                    timeout=self.timeout
                )
                response.raise_for_status()
                data = response.json()
                embeddings.append(data["embedding"])

            except requests.exceptions.ConnectionError:
                logger.error(
                    f"âŒ Ollama baÄŸlantÄ± hatasÄ±. "
                    f"'ollama serve' Ã§alÄ±ÅŸÄ±yor mu? ({self.endpoint})"
                )
                # SÄ±fÄ±r vektÃ¶r dÃ¶nerek ChromaDB'nin Ã§Ã¶kmesini Ã¶nle
                embeddings.append([0.0] * 768)

            except requests.exceptions.HTTPError as e:
                if e.response.status_code == 404:
                    logger.error(
                        f"âŒ Ollama embedding modeli bulunamadÄ±: {self.model}\n"
                        f"   Ã‡Ã¶zÃ¼m: ollama pull {self.model}"
                    )
                else:
                    logger.error(f"âŒ Ollama HTTP hatasÄ±: {e}")
                embeddings.append([0.0] * 768)

            except (KeyError, Exception) as e:
                logger.error(f"âŒ Embedding Ã¼retme hatasÄ±: {e}")
                embeddings.append([0.0] * 768)

        return embeddings

    def embed_query(self, input) -> List[List[float]]:
        """
        ChromaDB 0.4.x embed_query â€” __call__ arayÃ¼zÃ¼ne devreder.

        ChromaDB bu metodu 'input' keyword arg ile Ã§aÄŸÄ±rÄ±r; input bir str
        ya da List[str] olabilir. List[List[float]] dÃ¶ndÃ¼rÃ¼r (embedding matrisi).
        YanlÄ±ÅŸ sarmalama yapÄ±lmadan doÄŸrudan __call__'a iletilir;
        bÃ¶ylece Ollama'ya yanlÄ±ÅŸ formatta istek gÃ¶nderilmesi Ã¶nlenir.
        """
        if isinstance(input, list):
            return self(input)
        return self([input])

    def embed_documents(self, input) -> List[List[float]]:
        """
        ChromaDB 0.4.x embed_documents â€” __call__ arayÃ¼zÃ¼ne devreder.

        ChromaDB bu metodu 'input' keyword arg ile Ã§aÄŸÄ±rÄ±r; input bir str
        ya da List[str] olabilir. List[List[float]] dÃ¶ndÃ¼rÃ¼r (embedding matrisi).
        """
        if isinstance(input, list):
            return self(input)
        return self([input])


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SABITLER
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class MemoryConfig:
    """Memory system konfigÃ¼rasyonu"""
    # Database
    DB_TIMEOUT = 10  # saniye
    DB_VERSION = 1
    
    # Chunking
    CHUNK_SIZE = 800  # karakter
    CHUNK_OVERLAP = 150
    MIN_CHUNK_LENGTH = 10
    
    # Retrieval
    MAX_RECENT_MESSAGES = 10
    MAX_HISTORY_RESULTS = 3
    MAX_DOCUMENT_RESULTS = 4
    MIN_MESSAGE_LENGTH_FOR_VECTOR = 15
    
    # Ingestion
    SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.md', '.docx']
    
    # Embedding â€” Gemini/offline modu iÃ§in SentenceTransformer modeli
    EMBEDDING_MODEL = "paraphrase-multilingual-MiniLM-L12-v2"
    # Alternatif (hafif): "all-MiniLM-L6-v2"
    # Ollama modu iÃ§in Config.LOCAL_VEK kullanÄ±lÄ±r (nomic-embed-text)


class MessageRole(Enum):
    """Mesaj rolleri"""
    USER = "user"
    ASSISTANT = "assistant"
    MODEL = "model"
    SYSTEM = "system"


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# VERÄ° YAPILARI
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
@dataclass
class ChatMessage:
    """Chat mesaj yapÄ±sÄ±"""
    role: str
    content: str
    agent: Optional[str] = None
    timestamp: Optional[datetime] = None


@dataclass
class ProcessedFile:
    """Ä°ÅŸlenmiÅŸ dosya bilgisi"""
    filename: str
    file_hash: str
    file_size: int
    processed_date: datetime


@dataclass
class DocumentChunk:
    """Belge parÃ§asÄ±"""
    content: str
    source: str
    chunk_id: int
    file_hash: str
    file_type: str


@dataclass
class MemoryContext:
    """HafÄ±za baÄŸlamÄ±"""
    recent_messages: List[Dict[str, str]]
    long_term_history: str
    relevant_documents: str


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# MEMORY MANAGER
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
class MemoryManager:
    """
    LotusAI Hybrid Memory System
    
    Mimari:
    â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”
    â”‚          Memory Manager                     â”‚
    â”œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¬â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¤
    â”‚   SQLite        â”‚   ChromaDB (Vector DB)    â”‚
    â”‚   (Fast Access) â”‚   (Semantic Search)       â”‚
    â”œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¼â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¤
    â”‚ â€¢ Chat History  â”‚ â€¢ Chat Embeddings         â”‚
    â”‚ â€¢ File Tracking â”‚ â€¢ Document Embeddings     â”‚
    â”‚                 â”‚ â€¢ GPU Accelerated         â”‚
    â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”´â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜
    
    Features:
    - Thread-safe operations
    - GPU-accelerated embeddings
    - RAG (Retrieval-Augmented Generation)
    - Multi-format document support
    - Smart text chunking
    - Deduplication (hash-based)

    Embedding Stratejisi:
    - AI_PROVIDER='ollama' â†’ OllamaEmbeddingFunction (Config.LOCAL_VEK)
    - AI_PROVIDER='gemini' â†’ SentenceTransformerEmbeddingFunction (GPU destekli)
    """
    
    def __init__(self):
        """Memory manager baÅŸlatÄ±cÄ±"""
        # Paths
        self.work_dir = Config.WORK_DIR
        self.db_path = self.work_dir / "lotus_system.db"
        self.docs_path = self.work_dir / "documents"
        self.vector_db_path = self.work_dir / "lotus_vector_db"
        
        # Thread safety
        self.lock = threading.RLock()  # Reentrant lock
        
        # Create directories
        self.docs_path.mkdir(parents=True, exist_ok=True)
        self.vector_db_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize SQLite
        self._init_sqlite()
        
        # ChromaDB components
        self.chroma_client: Optional[chromadb.Client] = None
        self.history_collection = None
        self.docs_collection = None
        self.embedding_fn = None
        
        # Initialize ChromaDB
        if CHROMA_AVAILABLE:
            self._init_chromadb()
        
        logger.info("âœ… Memory Manager baÅŸlatÄ±ldÄ±")
    
    def _init_sqlite(self) -> None:
        """SQLite veritabanÄ±nÄ± baÅŸlat"""
        with self.lock:
            try:
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    
                    # Chat history table
                    cursor.execute("""
                        CREATE TABLE IF NOT EXISTS chat_history (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            agent TEXT NOT NULL,
                            role TEXT NOT NULL,
                            message TEXT NOT NULL,
                            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                        )
                    """)
                    
                    # Processed files table
                    cursor.execute("""
                        CREATE TABLE IF NOT EXISTS processed_files (
                            filename TEXT PRIMARY KEY,
                            processed_date DATETIME NOT NULL,
                            file_hash TEXT NOT NULL,
                            file_size INTEGER NOT NULL
                        )
                    """)
                    
                    # Schema version table
                    cursor.execute("""
                        CREATE TABLE IF NOT EXISTS schema_version (
                            version INTEGER PRIMARY KEY,
                            applied_date DATETIME DEFAULT CURRENT_TIMESTAMP
                        )
                    """)
                    
                    # Insert schema version if not exists
                    cursor.execute(
                        "INSERT OR IGNORE INTO schema_version (version) VALUES (?)",
                        (MemoryConfig.DB_VERSION,)
                    )
                    
                    # Indexes for performance
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_chat_agent "
                        "ON chat_history(agent)"
                    )
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_chat_timestamp "
                        "ON chat_history(timestamp DESC)"
                    )
                    
                    conn.commit()
                
                logger.info("âœ… SQLite veritabanÄ± hazÄ±r")
            
            except sqlite3.Error as e:
                logger.error(f"âŒ SQLite baÅŸlatma hatasÄ±: {e}")
                raise
    
    def _init_chromadb(self) -> None:
        """
        ChromaDB ve vektÃ¶r koleksiyonlarÄ±nÄ± baÅŸlat.

        AI_PROVIDER'a gÃ¶re embedding fonksiyonu seÃ§ilir:
        - 'ollama' â†’ OllamaEmbeddingFunction (Config.LOCAL_VEK)
        - 'gemini' â†’ SentenceTransformerEmbeddingFunction (GPU destekli)
        """
        try:
            # â”€â”€ Embedding fonksiyonu seÃ§imi â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            if Config.AI_PROVIDER == "ollama":
                self.embedding_fn = OllamaEmbeddingFunction(
                    model=Config.LOCAL_VEK,
                    ollama_url=Config.OLLAMA_URL,
                    timeout=Config.OLLAMA_TIMEOUT
                )
                logger.info(
                    f"ğŸ”¢ Embedding: Ollama ({Config.LOCAL_VEK})"
                )

            else:
                # Gemini / offline modu â€” SentenceTransformer (GPU)
                if Config.USE_GPU and torch.cuda.is_available():
                    device = "cuda"
                    logger.info(f"ğŸš€ GPU aktif: {torch.cuda.get_device_name(0)}")
                else:
                    device = "cpu"
                    if Config.USE_GPU:
                        logger.warning("âš ï¸ GPU istendi ama CUDA yok, CPU kullanÄ±lÄ±yor")

                self.embedding_fn = (
                    embedding_functions.SentenceTransformerEmbeddingFunction(
                        model_name=MemoryConfig.EMBEDDING_MODEL,
                        device=device
                    )
                )
                logger.info(
                    f"ğŸ”¢ Embedding: SentenceTransformer "
                    f"({MemoryConfig.EMBEDDING_MODEL} / {device.upper()})"
                )
            
            # ChromaDB client
            self.chroma_client = chromadb.PersistentClient(
                path=str(self.vector_db_path)
            )
            
            # Collections
            self.history_collection = self.chroma_client.get_or_create_collection(
                name="chat_history",
                metadata={"hnsw:space": "cosine"},
                embedding_function=self.embedding_fn
            )
            
            self.docs_collection = self.chroma_client.get_or_create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"},
                embedding_function=self.embedding_fn
            )
            
            doc_count = self.docs_collection.count()
            logger.info(f"âœ… ChromaDB aktif: {doc_count} belge parÃ§asÄ±")
        
        except Exception as e:
            logger.error(f"âŒ ChromaDB baÅŸlatma hatasÄ±: {e}")
            self.chroma_client = None
    
    @contextmanager
    def _get_db_connection(self):
        """
        Database connection context manager
        
        Yields:
            sqlite3.Connection
        """
        conn = None
        try:
            conn = sqlite3.connect(
                str(self.db_path),
                timeout=MemoryConfig.DB_TIMEOUT,
                check_same_thread=False
            )
            conn.row_factory = sqlite3.Row
            yield conn
        finally:
            if conn:
                conn.close()
    
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # DOCUMENT INGESTION (RAG)
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
    def ingest_documents(self) -> str:
        """
        Belgeleri tara ve vektÃ¶r veritabanÄ±na indeksle
        
        Returns:
            Ä°ÅŸlem sonucu mesajÄ±
        """
        if not CHROMA_AVAILABLE or not self.docs_collection:
            return "âŒ VektÃ¶r veritabanÄ± aktif deÄŸil"
        
        new_files = 0
        
        for ext in MemoryConfig.SUPPORTED_EXTENSIONS:
            pattern = f"*{ext}"
            
            for file_path in self.docs_path.glob(pattern):
                try:
                    if self._process_document(file_path):
                        new_files += 1
                
                except Exception as e:
                    logger.error(f"Dosya iÅŸleme hatasÄ± ({file_path.name}): {e}")
        
        if new_files > 0:
            return f"âœ… {new_files} yeni belge indekslendi"
        
        return "â„¹ï¸ Yeni belge bulunamadÄ±"
    
    def _process_document(self, file_path: Path) -> bool:
        """
        Tek bir belgeyi iÅŸle
        
        Args:
            file_path: Dosya yolu
        
        Returns:
            Ä°ÅŸlem baÅŸarÄ±lÄ±ysa True
        """
        filename = file_path.name
        
        # File hash ve metadata
        file_content_raw = file_path.read_bytes()
        file_hash = hashlib.md5(file_content_raw).hexdigest()
        file_size = file_path.stat().st_size
        
        # Daha Ã¶nce iÅŸlendi mi?
        if self._is_file_processed(filename, file_hash):
            return False
        
        # Ä°Ã§eriÄŸi oku
        content = self._read_document_content(file_path)
        
        if not content or not content.strip():
            logger.warning(f"BoÅŸ iÃ§erik: {filename}")
            return False
        
        # Chunk'la
        chunks = self._smart_chunk_text(content)
        
        if not chunks:
            logger.warning(f"Chunk oluÅŸturulamadÄ±: {filename}")
            return False
        
        # VektÃ¶r DB'ye ekle
        ids = [
            f"{filename}_{i}_{file_hash[:8]}"
            for i in range(len(chunks))
        ]
        
        metadatas = [
            {
                "source": filename,
                "chunk_id": i,
                "date": datetime.now().isoformat(),
                "type": file_path.suffix[1:],
                "hash": file_hash
            }
            for i in range(len(chunks))
        ]
        
        self.docs_collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        
        # Ä°ÅŸlenmiÅŸ olarak iÅŸaretle
        self._mark_file_as_processed(filename, file_hash, file_size)
        
        logger.info(f"ğŸ“„ Ä°ndekslendi: {filename} ({len(chunks)} chunk)")
        return True
    
    def _read_document_content(self, file_path: Path) -> str:
        """
        Dosya tipine gÃ¶re iÃ§eriÄŸi oku
        
        Args:
            file_path: Dosya yolu
        
        Returns:
            Dosya iÃ§eriÄŸi
        """
        ext = file_path.suffix.lower()
        
        # Text/Markdown
        if ext in ['.txt', '.md']:
            return file_path.read_text(encoding='utf-8', errors='ignore')
        
        # PDF
        if ext == '.pdf' and PDF_AVAILABLE:
            content = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            return "\n".join(content)
        
        # DOCX
        if ext == '.docx' and DOCX_AVAILABLE:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        
        return ""
    
    def _smart_chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None
    ) -> List[str]:
        """
        Metni anlamsal bÃ¼tÃ¼nlÃ¼ÄŸÃ¼ koruyarak chunk'la
        
        Args:
            text: Ä°ÅŸlenecek metin
            chunk_size: Chunk boyutu (karakter)
            overlap: Ã–rtÃ¼ÅŸme miktarÄ±
        
        Returns:
            Chunk listesi
        """
        if not text:
            return []
        
        chunk_size = chunk_size or MemoryConfig.CHUNK_SIZE
        overlap = overlap or MemoryConfig.CHUNK_OVERLAP
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = start + chunk_size
            
            # CÃ¼mle sonunda bitir (mÃ¼mkÃ¼nse)
            if end < text_len:
                # Ã–nce chunk'Ä±n ortasÄ±ndan sonra ara
                search_start = start + (chunk_size // 2)
                
                best_break = -1
                for punct in ['.', '!', '?', '\n']:
                    pos = text.rfind(punct, search_start, end)
                    if pos > best_break:
                        best_break = pos
                
                if best_break != -1:
                    end = best_break + 1
            
            chunk = text[start:end].strip()
            
            if len(chunk) >= MemoryConfig.MIN_CHUNK_LENGTH:
                chunks.append(chunk)
            
            # Overlap ile devam et
            start = max(start + 1, end - overlap)
            
            if end >= text_len:
                break
        
        return chunks
    
    def _is_file_processed(self, filename: str, file_hash: str) -> bool:
        """
        DosyanÄ±n iÅŸlenip iÅŸlenmediÄŸini kontrol et
        
        Args:
            filename: Dosya adÄ±
            file_hash: Dosya hash'i
        
        Returns:
            Ä°ÅŸlendiyse True
        """
        with self.lock:
            try:
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT 1 FROM processed_files "
                        "WHERE filename = ? AND file_hash = ?",
                        (filename, file_hash)
                    )
                    return cursor.fetchone() is not None
            
            except Exception as e:
                logger.error(f"Dosya kontrolÃ¼ hatasÄ±: {e}")
                return False
    
    def _mark_file_as_processed(
        self,
        filename: str,
        file_hash: str,
        file_size: int
    ) -> None:
        """
        DosyayÄ± iÅŸlenmiÅŸ olarak iÅŸaretle
        
        Args:
            filename: Dosya adÄ±
            file_hash: Dosya hash'i
            file_size: Dosya boyutu
        """
        with self.lock:
            try:
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        """
                        INSERT OR REPLACE INTO processed_files
                        (filename, processed_date, file_hash, file_size)
                        VALUES (?, ?, ?, ?)
                        """,
                        (filename, datetime.now().isoformat(), file_hash, file_size)
                    )
                    conn.commit()
            
            except Exception as e:
                logger.error(f"Dosya iÅŸaretleme hatasÄ±: {e}")
    
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # CHAT MEMORY
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
    def save(
        self,
        agent: str,
        role: str,
        message: str
    ) -> None:
        """
        MesajÄ± hafÄ±zaya kaydet
        
        Args:
            agent: Agent adÄ±
            role: Mesaj rolÃ¼ (user/assistant)
            message: Mesaj iÃ§eriÄŸi
        """
        if not message or not message.strip():
            return
        
        timestamp = datetime.now()
        
        # 1. SQLite'a kaydet
        with self.lock:
            try:
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        """
                        INSERT INTO chat_history (agent, role, message, timestamp)
                        VALUES (?, ?, ?, ?)
                        """,
                        (agent, role, message, timestamp.isoformat())
                    )
                    conn.commit()
            
            except sqlite3.Error as e:
                logger.error(f"Mesaj kaydetme hatasÄ±: {e}")
        
        # 2. ChromaDB'ye vektÃ¶r olarak kaydet
        if (CHROMA_AVAILABLE and
            self.history_collection and
            len(message) >= MemoryConfig.MIN_MESSAGE_LENGTH_FOR_VECTOR):
            
            try:
                unique_id = (
                    f"{agent}_{timestamp.timestamp()}_"
                    f"{hashlib.md5(message.encode()).hexdigest()[:6]}"
                )
                
                metadata = {
                    "agent": agent,
                    "role": role,
                    "time": timestamp.isoformat()
                }
                
                self.history_collection.add(
                    documents=[message],
                    metadatas=[metadata],
                    ids=[unique_id]
                )
            
            except Exception as e:
                logger.debug(f"VektÃ¶r kayÄ±t hatasÄ±: {e}")
    
    def load_context(
        self,
        agent: str,
        query: str,
        max_items: Optional[int] = None,
        include_semantic: bool = True
    ) -> Tuple[List[Dict[str, str]], str, str]:
        """
        Agent iÃ§in baÄŸlam yÃ¼kle

        Args:
            agent: Agent adÄ±
            query: Sorgu metni
            max_items: Maksimum mesaj sayÄ±sÄ±
            include_semantic: ChromaDB semantik aramasÄ± yapÄ±lsÄ±n mÄ± (varsayÄ±lan True).
                              Sadece son mesajlar gerekiyorsa False geÃ§ilerek
                              gereksiz vektÃ¶r aramasÄ± atlanÄ±r.

        Returns:
            Tuple[son mesajlar, uzun dÃ¶nem hafÄ±za, ilgili dokÃ¼manlar]
        """
        max_items = max_items or MemoryConfig.MAX_RECENT_MESSAGES

        # 1. Son mesajlar (SQLite â€” her zaman)
        recent_messages = self._get_recent_messages(agent, max_items)

        if not include_semantic:
            return recent_messages, "", ""

        # 2. Uzun dÃ¶nem hafÄ±za (ChromaDB - semantic search)
        long_term_history = self._search_history(agent, query)

        # 3. Ä°lgili dokÃ¼manlar (RAG)
        relevant_docs = self._search_documents(query)

        return recent_messages, long_term_history, relevant_docs
    
    def _get_recent_messages(
        self,
        agent: str,
        limit: int
    ) -> List[Dict[str, str]]:
        """
        Son mesajlarÄ± getir
        
        Args:
            agent: Agent adÄ±
            limit: Maksimum sayÄ±
        
        Returns:
            Mesaj listesi
        """
        with self.lock:
            try:
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        """
                        SELECT role, message
                        FROM chat_history
                        WHERE agent = ? OR agent = 'SYSTEM'
                        ORDER BY id DESC
                        LIMIT ?
                        """,
                        (agent, limit)
                    )
                    rows = cursor.fetchall()
                    
                    # Reverse to get chronological order
                    return [
                        {"role": row["role"], "content": row["message"]}
                        for row in reversed(rows)
                    ]
            
            except Exception as e:
                logger.error(f"Son mesajlar getirme hatasÄ±: {e}")
                return []
    
    def _search_history(self, agent: str, query: str) -> str:
        """
        GeÃ§miÅŸ sohbetlerde anlamsal arama
        
        Args:
            agent: Agent adÄ±
            query: Arama sorgusu
        
        Returns:
            Ä°lgili geÃ§miÅŸ mesajlar
        """
        if not CHROMA_AVAILABLE or not self.history_collection:
            return ""
        
        try:
            results = self.history_collection.query(
                query_texts=[query],
                n_results=MemoryConfig.MAX_HISTORY_RESULTS,
                where={"agent": agent}
            )
            
            if results['documents'] and results['documents'][0]:
                return "\n".join([
                    f"â€¢ {doc}"
                    for doc in results['documents'][0]
                ])
        
        except Exception as e:
            logger.error(f"GeÃ§miÅŸ arama hatasÄ±: {e}")
        
        return ""
    
    def _search_documents(self, query: str) -> str:
        """
        Belgelerde anlamsal arama (RAG)
        
        Args:
            query: Arama sorgusu
        
        Returns:
            Ä°lgili belge parÃ§alarÄ±
        """
        if not CHROMA_AVAILABLE or not self.docs_collection:
            return ""
        
        try:
            results = self.docs_collection.query(
                query_texts=[query],
                n_results=MemoryConfig.MAX_DOCUMENT_RESULTS
            )
            
            if results['documents'] and results['documents'][0]:
                docs_content = []
                
                for i, doc_text in enumerate(results['documents'][0]):
                    source = results['metadatas'][0][i].get('source', 'Bilinmiyor')
                    docs_content.append(f"[{source}]: {doc_text}")
                
                return "\n\n".join(docs_content)
        
        except Exception as e:
            logger.error(f"Belge arama hatasÄ±: {e}")
        
        return ""
    
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # WEB INTERFACE
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
    def get_agent_history_for_web(
        self,
        agent_name: str,
        limit: int = 40
    ) -> List[Dict[str, Any]]:
        """
        Web arayÃ¼zÃ¼ iÃ§in geÃ§miÅŸ getir
        
        Args:
            agent_name: Agent adÄ±
            limit: Maksimum mesaj sayÄ±sÄ±
        
        Returns:
            FormatlanmÄ±ÅŸ mesaj listesi
        """
        target_agent = "ATLAS" if agent_name == "GENEL" else agent_name
        
        with self.lock:
            try:
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        """
                        SELECT agent, role, message, timestamp
                        FROM chat_history
                        WHERE agent = ?
                        ORDER BY id DESC
                        LIMIT ?
                        """,
                        (target_agent, limit)
                    )
                    rows = cursor.fetchall()
                    
                    return [
                        {
                            "sender": "Siz" if row["role"] == "user" else row["agent"],
                            "text": row["message"],
                            "type": "user" if row["role"] == "user" else "agent",
                            "time": row["timestamp"][11:16] if row["timestamp"] else ""
                        }
                        for row in reversed(rows)
                    ]
            
            except Exception as e:
                logger.error(f"Web geÃ§miÅŸi hatasÄ±: {e}")
                return []
    
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # CLEANUP & MANAGEMENT
    # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
    def clear_history(self, only_chat: bool = True) -> bool:
        """
        HafÄ±zayÄ± temizle
        
        Args:
            only_chat: Sadece sohbet geÃ§miÅŸi mi (True) yoksa tÃ¼m hafÄ±za mÄ± (False)
        
        Returns:
            BaÅŸarÄ±lÄ±ysa True
        """
        with self.lock:
            try:
                # SQLite temizliÄŸi
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute("DELETE FROM chat_history")
                    
                    if not only_chat:
                        cursor.execute("DELETE FROM processed_files")
                    
                    conn.commit()
                
                # ChromaDB temizliÄŸi
                if CHROMA_AVAILABLE and self.chroma_client:
                    try:
                        # Chat history collection
                        self.chroma_client.delete_collection("chat_history")
                        self.history_collection = self.chroma_client.get_or_create_collection(
                            name="chat_history",
                            metadata={"hnsw:space": "cosine"},
                            embedding_function=self.embedding_fn
                        )
                        
                        # Documents collection
                        if not only_chat:
                            self.chroma_client.delete_collection("documents")
                            self.docs_collection = self.chroma_client.get_or_create_collection(
                                name="documents",
                                metadata={"hnsw:space": "cosine"},
                                embedding_function=self.embedding_fn
                            )
                    
                    except Exception as e:
                        logger.warning(f"ChromaDB temizleme uyarÄ±sÄ±: {e}")
                
                mode = "Sohbet" if only_chat else "Tam"
                logger.info(f"ğŸ§¹ HafÄ±za temizlendi ({mode})")
                return True
            
            except Exception as e:
                logger.error(f"HafÄ±za temizleme hatasÄ±: {e}")
                return False
    
    def get_stats(self) -> Dict[str, Any]:
        """
        HafÄ±za istatistiklerini getir
        
        Returns:
            Ä°statistik dictionary'si
        """
        stats = {
            "total_messages": 0,
            "total_documents": 0,
            "processed_files": 0,
            "db_size_mb": 0,
            "vector_db_enabled": CHROMA_AVAILABLE,
            "embedding_backend": (
                f"Ollama ({Config.LOCAL_VEK})"
                if Config.AI_PROVIDER == "ollama"
                else f"SentenceTransformer ({MemoryConfig.EMBEDDING_MODEL})"
            )
        }
        
        with self.lock:
            try:
                # SQLite stats
                with self._get_db_connection() as conn:
                    cursor = conn.cursor()
                    
                    # Message count
                    cursor.execute("SELECT COUNT(*) FROM chat_history")
                    stats["total_messages"] = cursor.fetchone()[0]
                    
                    # Processed files
                    cursor.execute("SELECT COUNT(*) FROM processed_files")
                    stats["processed_files"] = cursor.fetchone()[0]
                
                # DB file size
                if self.db_path.exists():
                    stats["db_size_mb"] = round(
                        self.db_path.stat().st_size / (1024 * 1024),
                        2
                    )
                
                # ChromaDB stats
                if CHROMA_AVAILABLE and self.docs_collection:
                    stats["total_documents"] = self.docs_collection.count()
            
            except Exception as e:
                logger.error(f"Ä°statistik hatasÄ±: {e}")
        
        return stats
    
    def shutdown(self) -> None:
        """Memory manager'Ä± kapat"""
        logger.info("Memory Manager kapatÄ±lÄ±yor...")
        
        # ChromaDB'yi kapat (gerekirse)
        if self.chroma_client:
            try:
                # ChromaDB otomatik olarak persist ediyor
                self.chroma_client = None
            except Exception as e:
                logger.warning(f"ChromaDB kapatma uyarÄ±sÄ±: {e}")
        
        logger.info("âœ… Memory Manager kapatÄ±ldÄ±")

# """
# LotusAI Memory Management System
# SÃ¼rÃ¼m: 2.5.4
# AÃ§Ä±klama: Hybrid hafÄ±za sistemi (SQLite + ChromaDB) ve RAG implementasyonu

# Ã–zellikler:
# - KÄ±sa sÃ¼reli hafÄ±za: SQLite (hÄ±zlÄ± eriÅŸim)
# - Uzun sÃ¼reli hafÄ±za: ChromaDB (anlamsal arama)
# - RAG: Belge indeksleme ve retrieval
# - GPU hÄ±zlandÄ±rma
# - Thread-safe operasyonlar
# - Ollama embedding desteÄŸi (nomic-embed-text) â€” AI_PROVIDER='ollama' iÃ§in
# """

# import sqlite3
# import logging
# import threading
# import hashlib
# import re
# import requests
# from pathlib import Path
# from datetime import datetime
# from typing import List, Dict, Tuple, Any, Optional
# from dataclasses import dataclass
# from contextlib import contextmanager
# from enum import Enum

# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # CONFIG
# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# from config import Config

# logger = logging.getLogger("LotusAI.Memory")


# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # EXTERNAL LIBRARIES
# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # PDF
# try:
#     import PyPDF2
#     PDF_AVAILABLE = True
# except ImportError:
#     PDF_AVAILABLE = False
#     logger.warning("âš ï¸ PyPDF2 yÃ¼klÃ¼ deÄŸil, PDF desteÄŸi yok")

# # DOCX
# try:
#     import docx
#     DOCX_AVAILABLE = True
# except ImportError:
#     DOCX_AVAILABLE = False
#     logger.warning("âš ï¸ python-docx yÃ¼klÃ¼ deÄŸil, DOCX desteÄŸi yok")

# # ChromaDB + Torch
# try:
#     import chromadb
#     from chromadb.config import Settings
#     from chromadb.utils import embedding_functions
#     import torch
#     CHROMA_AVAILABLE = True
# except ImportError:
#     CHROMA_AVAILABLE = False
#     logger.warning("âš ï¸ ChromaDB/torch yÃ¼klÃ¼ deÄŸil, vektÃ¶r arama devre dÄ±ÅŸÄ±")


# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # OLLAMA EMBEDDING FUNCTION
# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# class OllamaEmbeddingFunction:
#     """
#     ChromaDB uyumlu Ollama embedding fonksiyonu.

#     Config.LOCAL_VEK modelini (varsayÄ±lan: nomic-embed-text) kullanarak
#     Ollama'nÄ±n /api/embeddings endpoint'inden vektÃ¶r Ã¼retir.

#     KullanÄ±m: AI_PROVIDER='ollama' olduÄŸunda SentenceTransformer'Ä±n yerini alÄ±r.
#     """

#     def __init__(self, model: str, ollama_url: str, timeout: int = 30):
#         """
#         Args:
#             model: Ollama embedding model adÄ± (Ã¶rn. nomic-embed-text)
#             ollama_url: Ollama base URL (Ã¶rn. http://localhost:11434/api)
#             timeout: Ä°stek zaman aÅŸÄ±mÄ± (saniye)
#         """
#         self.model = model
#         self.timeout = timeout

#         # /api/embeddings endpoint'ini doÄŸru ÅŸekilde oluÅŸtur
#         base = ollama_url.rstrip("/")
#         if base.endswith("/api"):
#             self.endpoint = f"{base}/embeddings"
#         elif "/api" in base:
#             # Zaten /api/chat gibi bir ÅŸey varsa base'i al
#             self.endpoint = base.rsplit("/", 1)[0] + "/embeddings"
#         else:
#             self.endpoint = f"{base}/api/embeddings"

#         logger.info(
#             f"ğŸ”¢ OllamaEmbeddingFunction hazÄ±r | "
#             f"Model: {self.model} | Endpoint: {self.endpoint}"
#         )

#     def __call__(self, input: List[str]) -> List[List[float]]:
#         """
#         ChromaDB'nin beklediÄŸi arayÃ¼z.

#         Args:
#             input: VektÃ¶rlenecek metin listesi

#         Returns:
#             Embedding vektÃ¶rleri listesi
#         """
#         embeddings = []

#         for text in input:
#             try:
#                 response = requests.post(
#                     self.endpoint,
#                     json={"model": self.model, "prompt": text},
#                     timeout=self.timeout
#                 )
#                 response.raise_for_status()
#                 data = response.json()
#                 embeddings.append(data["embedding"])

#             except requests.exceptions.ConnectionError:
#                 logger.error(
#                     f"âŒ Ollama baÄŸlantÄ± hatasÄ±. "
#                     f"'ollama serve' Ã§alÄ±ÅŸÄ±yor mu? ({self.endpoint})"
#                 )
#                 # SÄ±fÄ±r vektÃ¶r dÃ¶nerek ChromaDB'nin Ã§Ã¶kmesini Ã¶nle
#                 embeddings.append([0.0] * 768)

#             except requests.exceptions.HTTPError as e:
#                 if e.response.status_code == 404:
#                     logger.error(
#                         f"âŒ Ollama embedding modeli bulunamadÄ±: {self.model}\n"
#                         f"   Ã‡Ã¶zÃ¼m: ollama pull {self.model}"
#                     )
#                 else:
#                     logger.error(f"âŒ Ollama HTTP hatasÄ±: {e}")
#                 embeddings.append([0.0] * 768)

#             except (KeyError, Exception) as e:
#                 logger.error(f"âŒ Embedding Ã¼retme hatasÄ±: {e}")
#                 embeddings.append([0.0] * 768)

#         return embeddings


# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # SABITLER
# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# class MemoryConfig:
#     """Memory system konfigÃ¼rasyonu"""
#     # Database
#     DB_TIMEOUT = 10
#     DB_VERSION = 1

#     # Chunking
#     CHUNK_SIZE = 800
#     CHUNK_OVERLAP = 150
#     MIN_CHUNK_LENGTH = 10

#     # Retrieval
#     MAX_RECENT_MESSAGES = 10
#     MAX_HISTORY_RESULTS = 3
#     MAX_DOCUMENT_RESULTS = 4
#     MIN_MESSAGE_LENGTH_FOR_VECTOR = 15

#     # Ingestion
#     SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.md', '.docx']

#     # Embedding â€” Gemini/offline modu iÃ§in SentenceTransformer modeli
#     SENTENCE_TRANSFORMER_MODEL = "paraphrase-multilingual-MiniLM-L12-v2"
#     # Ollama modu iÃ§in Config.LOCAL_VEK kullanÄ±lÄ±r (nomic-embed-text)


# class MessageRole(Enum):
#     USER = "user"
#     ASSISTANT = "assistant"
#     MODEL = "model"
#     SYSTEM = "system"


# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # VERÄ° YAPILARI
# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# @dataclass
# class ChatMessage:
#     role: str
#     content: str
#     agent: Optional[str] = None
#     timestamp: Optional[datetime] = None


# @dataclass
# class ProcessedFile:
#     filename: str
#     file_hash: str
#     file_size: int
#     processed_date: datetime


# @dataclass
# class DocumentChunk:
#     content: str
#     source: str
#     chunk_id: int
#     file_hash: str
#     file_type: str


# @dataclass
# class MemoryContext:
#     recent_messages: List[Dict[str, str]]
#     long_term_history: str
#     relevant_documents: str


# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # MEMORY MANAGER
# # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# class MemoryManager:
#     """
#     LotusAI Hybrid Memory System

#     Mimari:
#     â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”
#     â”‚                    Memory Manager                        â”‚
#     â”œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¬â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¤
#     â”‚   SQLite             â”‚   ChromaDB (Vector DB)            â”‚
#     â”‚   (Fast Access)      â”‚   (Semantic Search)               â”‚
#     â”œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¼â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¤
#     â”‚ â€¢ Chat History       â”‚ â€¢ Chat Embeddings                 â”‚
#     â”‚ â€¢ File Tracking      â”‚ â€¢ Document Embeddings             â”‚
#     â”‚                      â”‚ â€¢ GPU Accelerated (Gemini modu)   â”‚
#     â”‚                      â”‚ â€¢ Ollama Embed (Ollama modu)      â”‚
#     â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”´â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜

#     Embedding Stratejisi:
#     - AI_PROVIDER='ollama' â†’ OllamaEmbeddingFunction (Config.LOCAL_VEK)
#     - AI_PROVIDER='gemini' â†’ SentenceTransformerEmbeddingFunction (GPU)
#     """

#     def __init__(self):
#         # Paths
#         self.work_dir = Config.WORK_DIR
#         self.db_path = self.work_dir / "lotus_system.db"
#         self.docs_path = self.work_dir / "documents"
#         self.vector_db_path = self.work_dir / "lotus_vector_db"

#         # Thread safety
#         self.lock = threading.RLock()

#         # Create directories
#         self.docs_path.mkdir(parents=True, exist_ok=True)
#         self.vector_db_path.mkdir(parents=True, exist_ok=True)

#         # Initialize SQLite
#         self._init_sqlite()

#         # ChromaDB components
#         self.chroma_client: Optional[chromadb.Client] = None
#         self.history_collection = None
#         self.docs_collection = None
#         self.embedding_fn = None

#         # Initialize ChromaDB
#         if CHROMA_AVAILABLE:
#             self._init_chromadb()

#         logger.info("âœ… Memory Manager baÅŸlatÄ±ldÄ±")

#     def _init_sqlite(self) -> None:
#         """SQLite veritabanÄ±nÄ± baÅŸlat"""
#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()

#                     cursor.execute("""
#                         CREATE TABLE IF NOT EXISTS chat_history (
#                             id INTEGER PRIMARY KEY AUTOINCREMENT,
#                             agent TEXT NOT NULL,
#                             role TEXT NOT NULL,
#                             message TEXT NOT NULL,
#                             timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
#                         )
#                     """)

#                     cursor.execute("""
#                         CREATE TABLE IF NOT EXISTS processed_files (
#                             filename TEXT PRIMARY KEY,
#                             processed_date DATETIME NOT NULL,
#                             file_hash TEXT NOT NULL,
#                             file_size INTEGER NOT NULL
#                         )
#                     """)

#                     cursor.execute("""
#                         CREATE TABLE IF NOT EXISTS schema_version (
#                             version INTEGER PRIMARY KEY,
#                             applied_date DATETIME DEFAULT CURRENT_TIMESTAMP
#                         )
#                     """)

#                     cursor.execute(
#                         "INSERT OR IGNORE INTO schema_version (version) VALUES (?)",
#                         (MemoryConfig.DB_VERSION,)
#                     )

#                     cursor.execute(
#                         "CREATE INDEX IF NOT EXISTS idx_chat_agent "
#                         "ON chat_history(agent)"
#                     )
#                     cursor.execute(
#                         "CREATE INDEX IF NOT EXISTS idx_chat_timestamp "
#                         "ON chat_history(timestamp DESC)"
#                     )

#                     conn.commit()

#                 logger.info("âœ… SQLite veritabanÄ± hazÄ±r")

#             except sqlite3.Error as e:
#                 logger.error(f"âŒ SQLite baÅŸlatma hatasÄ±: {e}")
#                 raise

#     def _init_chromadb(self) -> None:
#         """
#         ChromaDB ve vektÃ¶r koleksiyonlarÄ±nÄ± baÅŸlat.

#         AI_PROVIDER'a gÃ¶re embedding fonksiyonu seÃ§ilir:
#         - 'ollama' â†’ OllamaEmbeddingFunction (Config.LOCAL_VEK)
#         - 'gemini' â†’ SentenceTransformerEmbeddingFunction (GPU destekli)
#         """
#         try:
#             # â”€â”€ Embedding fonksiyonu seÃ§imi â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#             if Config.AI_PROVIDER == "ollama":
#                 self.embedding_fn = OllamaEmbeddingFunction(
#                     model=Config.LOCAL_VEK,
#                     ollama_url=Config.OLLAMA_URL,
#                     timeout=Config.OLLAMA_TIMEOUT
#                 )
#                 logger.info(
#                     f"ğŸ”¢ Embedding: Ollama ({Config.LOCAL_VEK})"
#                 )

#             else:
#                 # Gemini / offline modu â€” SentenceTransformer (GPU)
#                 if Config.USE_GPU and torch.cuda.is_available():
#                     device = "cuda"
#                     logger.info(f"ğŸš€ Embedding GPU aktif: {torch.cuda.get_device_name(0)}")
#                 else:
#                     device = "cpu"
#                     if Config.USE_GPU:
#                         logger.warning("âš ï¸ GPU istendi ama CUDA yok, CPU kullanÄ±lÄ±yor")

#                 self.embedding_fn = (
#                     embedding_functions.SentenceTransformerEmbeddingFunction(
#                         model_name=MemoryConfig.SENTENCE_TRANSFORMER_MODEL,
#                         device=device
#                     )
#                 )
#                 logger.info(
#                     f"ğŸ”¢ Embedding: SentenceTransformer "
#                     f"({MemoryConfig.SENTENCE_TRANSFORMER_MODEL} / {device.upper()})"
#                 )

#             # â”€â”€ ChromaDB client â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#             self.chroma_client = chromadb.PersistentClient(
#                 path=str(self.vector_db_path)
#             )

#             # â”€â”€ Koleksiyonlar â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#             self.history_collection = self.chroma_client.get_or_create_collection(
#                 name="chat_history",
#                 metadata={"hnsw:space": "cosine"},
#                 embedding_function=self.embedding_fn
#             )

#             self.docs_collection = self.chroma_client.get_or_create_collection(
#                 name="documents",
#                 metadata={"hnsw:space": "cosine"},
#                 embedding_function=self.embedding_fn
#             )

#             doc_count = self.docs_collection.count()
#             logger.info(f"âœ… ChromaDB aktif: {doc_count} belge parÃ§asÄ±")

#         except Exception as e:
#             logger.error(f"âŒ ChromaDB baÅŸlatma hatasÄ±: {e}")
#             self.chroma_client = None

#     @contextmanager
#     def _get_db_connection(self):
#         """Database connection context manager"""
#         conn = None
#         try:
#             conn = sqlite3.connect(
#                 str(self.db_path),
#                 timeout=MemoryConfig.DB_TIMEOUT,
#                 check_same_thread=False
#             )
#             conn.row_factory = sqlite3.Row
#             yield conn
#         finally:
#             if conn:
#                 conn.close()

#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#     # DOCUMENT INGESTION (RAG)
#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

#     def ingest_documents(self) -> str:
#         """
#         Belgeleri tara ve vektÃ¶r veritabanÄ±na indeksle

#         Returns:
#             Ä°ÅŸlem sonucu mesajÄ±
#         """
#         if not CHROMA_AVAILABLE or not self.docs_collection:
#             return "âŒ VektÃ¶r veritabanÄ± aktif deÄŸil"

#         new_files = 0

#         for ext in MemoryConfig.SUPPORTED_EXTENSIONS:
#             for file_path in self.docs_path.glob(f"*{ext}"):
#                 try:
#                     if self._process_document(file_path):
#                         new_files += 1
#                 except Exception as e:
#                     logger.error(f"Dosya iÅŸleme hatasÄ± ({file_path.name}): {e}")

#         if new_files > 0:
#             return f"âœ… {new_files} yeni belge indekslendi"

#         return "â„¹ï¸ Yeni belge bulunamadÄ±"

#     def _process_document(self, file_path: Path) -> bool:
#         """
#         Tek bir belgeyi iÅŸle

#         Args:
#             file_path: Dosya yolu

#         Returns:
#             Ä°ÅŸlem baÅŸarÄ±lÄ±ysa True
#         """
#         filename = file_path.name
#         file_content_raw = file_path.read_bytes()
#         file_hash = hashlib.md5(file_content_raw).hexdigest()
#         file_size = file_path.stat().st_size

#         if self._is_file_processed(filename, file_hash):
#             return False

#         content = self._read_document_content(file_path)

#         if not content or not content.strip():
#             logger.warning(f"BoÅŸ iÃ§erik: {filename}")
#             return False

#         chunks = self._smart_chunk_text(content)

#         if not chunks:
#             logger.warning(f"Chunk oluÅŸturulamadÄ±: {filename}")
#             return False

#         ids = [
#             f"{filename}_{i}_{file_hash[:8]}"
#             for i in range(len(chunks))
#         ]

#         metadatas = [
#             {
#                 "source": filename,
#                 "chunk_id": i,
#                 "date": datetime.now().isoformat(),
#                 "type": file_path.suffix[1:],
#                 "hash": file_hash
#             }
#             for i in range(len(chunks))
#         ]

#         self.docs_collection.add(
#             documents=chunks,
#             metadatas=metadatas,
#             ids=ids
#         )

#         self._mark_file_as_processed(filename, file_hash, file_size)
#         logger.info(f"ğŸ“„ Ä°ndekslendi: {filename} ({len(chunks)} chunk)")
#         return True

#     def _read_document_content(self, file_path: Path) -> str:
#         """
#         Dosya tipine gÃ¶re iÃ§eriÄŸi oku

#         Args:
#             file_path: Dosya yolu

#         Returns:
#             Dosya iÃ§eriÄŸi
#         """
#         ext = file_path.suffix.lower()

#         if ext in ['.txt', '.md']:
#             return file_path.read_text(encoding='utf-8', errors='ignore')

#         if ext == '.pdf' and PDF_AVAILABLE:
#             content = []
#             with open(file_path, 'rb') as f:
#                 reader = PyPDF2.PdfReader(f)
#                 for page in reader.pages:
#                     text = page.extract_text()
#                     if text:
#                         content.append(text)
#             return "\n".join(content)

#         if ext == '.docx' and DOCX_AVAILABLE:
#             doc = docx.Document(file_path)
#             return "\n".join([para.text for para in doc.paragraphs])

#         return ""

#     def _smart_chunk_text(
#         self,
#         text: str,
#         chunk_size: Optional[int] = None,
#         overlap: Optional[int] = None
#     ) -> List[str]:
#         """
#         Metni anlamsal bÃ¼tÃ¼nlÃ¼ÄŸÃ¼ koruyarak chunk'la

#         Args:
#             text: Ä°ÅŸlenecek metin
#             chunk_size: Chunk boyutu (karakter)
#             overlap: Ã–rtÃ¼ÅŸme miktarÄ±

#         Returns:
#             Chunk listesi
#         """
#         if not text:
#             return []

#         chunk_size = chunk_size or MemoryConfig.CHUNK_SIZE
#         overlap = overlap or MemoryConfig.CHUNK_OVERLAP

#         text = re.sub(r'\s+', ' ', text).strip()

#         chunks = []
#         start = 0
#         text_len = len(text)

#         while start < text_len:
#             end = start + chunk_size

#             if end < text_len:
#                 search_start = start + (chunk_size // 2)
#                 best_break = -1

#                 for punct in ['.', '!', '?', '\n']:
#                     pos = text.rfind(punct, search_start, end)
#                     if pos > best_break:
#                         best_break = pos

#                 if best_break != -1:
#                     end = best_break + 1

#             chunk = text[start:end].strip()

#             if len(chunk) >= MemoryConfig.MIN_CHUNK_LENGTH:
#                 chunks.append(chunk)

#             start = max(start + 1, end - overlap)

#             if end >= text_len:
#                 break

#         return chunks

#     def _is_file_processed(self, filename: str, file_hash: str) -> bool:
#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute(
#                         "SELECT 1 FROM processed_files "
#                         "WHERE filename = ? AND file_hash = ?",
#                         (filename, file_hash)
#                     )
#                     return cursor.fetchone() is not None
#             except Exception as e:
#                 logger.error(f"Dosya kontrolÃ¼ hatasÄ±: {e}")
#                 return False

#     def _mark_file_as_processed(
#         self,
#         filename: str,
#         file_hash: str,
#         file_size: int
#     ) -> None:
#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute(
#                         """
#                         INSERT OR REPLACE INTO processed_files
#                         (filename, processed_date, file_hash, file_size)
#                         VALUES (?, ?, ?, ?)
#                         """,
#                         (filename, datetime.now().isoformat(), file_hash, file_size)
#                     )
#                     conn.commit()
#             except Exception as e:
#                 logger.error(f"Dosya iÅŸaretleme hatasÄ±: {e}")

#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#     # CHAT MEMORY
#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

#     def save(self, agent: str, role: str, message: str) -> None:
#         """
#         MesajÄ± hafÄ±zaya kaydet

#         Args:
#             agent: Agent adÄ±
#             role: Mesaj rolÃ¼ (user/assistant/model)
#             message: Mesaj iÃ§eriÄŸi
#         """
#         if not message or not message.strip():
#             return

#         timestamp = datetime.now()

#         # 1. SQLite
#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute(
#                         """
#                         INSERT INTO chat_history (agent, role, message, timestamp)
#                         VALUES (?, ?, ?, ?)
#                         """,
#                         (agent, role, message, timestamp.isoformat())
#                     )
#                     conn.commit()
#             except sqlite3.Error as e:
#                 logger.error(f"Mesaj kaydetme hatasÄ±: {e}")

#         # 2. ChromaDB vektÃ¶r kaydÄ±
#         if (CHROMA_AVAILABLE and
#                 self.history_collection and
#                 len(message) >= MemoryConfig.MIN_MESSAGE_LENGTH_FOR_VECTOR):
#             try:
#                 unique_id = (
#                     f"{agent}_{timestamp.timestamp()}_"
#                     f"{hashlib.md5(message.encode()).hexdigest()[:6]}"
#                 )

#                 self.history_collection.add(
#                     documents=[message],
#                     metadatas=[{
#                         "agent": agent,
#                         "role": role,
#                         "time": timestamp.isoformat()
#                     }],
#                     ids=[unique_id]
#                 )
#             except Exception as e:
#                 logger.debug(f"VektÃ¶r kayÄ±t hatasÄ±: {e}")

#     def load_context(
#         self,
#         agent: str,
#         query: str,
#         max_items: Optional[int] = None
#     ) -> Tuple[List[Dict[str, str]], str, str]:
#         """
#         Agent iÃ§in baÄŸlam yÃ¼kle

#         Args:
#             agent: Agent adÄ±
#             query: Sorgu metni
#             max_items: Maksimum mesaj sayÄ±sÄ±

#         Returns:
#             Tuple[son mesajlar, uzun dÃ¶nem hafÄ±za, ilgili dokÃ¼manlar]
#         """
#         max_items = max_items or MemoryConfig.MAX_RECENT_MESSAGES

#         recent_messages = self._get_recent_messages(agent, max_items)
#         long_term_history = self._search_history(agent, query)
#         relevant_docs = self._search_documents(query)

#         return recent_messages, long_term_history, relevant_docs

#     def _get_recent_messages(self, agent: str, limit: int) -> List[Dict[str, str]]:
#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute(
#                         """
#                         SELECT role, message
#                         FROM chat_history
#                         WHERE agent = ? OR agent = 'SYSTEM'
#                         ORDER BY id DESC
#                         LIMIT ?
#                         """,
#                         (agent, limit)
#                     )
#                     rows = cursor.fetchall()
#                     return [
#                         {"role": row["role"], "content": row["message"]}
#                         for row in reversed(rows)
#                     ]
#             except Exception as e:
#                 logger.error(f"Son mesajlar getirme hatasÄ±: {e}")
#                 return []

#     def _search_history(self, agent: str, query: str) -> str:
#         if not CHROMA_AVAILABLE or not self.history_collection:
#             return ""
#         try:
#             results = self.history_collection.query(
#                 query_texts=[query],
#                 n_results=MemoryConfig.MAX_HISTORY_RESULTS,
#                 where={"agent": agent}
#             )
#             if results['documents'] and results['documents'][0]:
#                 return "\n".join([f"â€¢ {doc}" for doc in results['documents'][0]])
#         except Exception as e:
#             logger.error(f"GeÃ§miÅŸ arama hatasÄ±: {e}")
#         return ""

#     def _search_documents(self, query: str) -> str:
#         if not CHROMA_AVAILABLE or not self.docs_collection:
#             return ""
#         try:
#             results = self.docs_collection.query(
#                 query_texts=[query],
#                 n_results=MemoryConfig.MAX_DOCUMENT_RESULTS
#             )
#             if results['documents'] and results['documents'][0]:
#                 docs_content = []
#                 for i, doc_text in enumerate(results['documents'][0]):
#                     source = results['metadatas'][0][i].get('source', 'Bilinmiyor')
#                     docs_content.append(f"[{source}]: {doc_text}")
#                 return "\n\n".join(docs_content)
#         except Exception as e:
#             logger.error(f"Belge arama hatasÄ±: {e}")
#         return ""

#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#     # WEB INTERFACE
#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

#     def get_agent_history_for_web(
#         self,
#         agent_name: str,
#         limit: int = 40
#     ) -> List[Dict[str, Any]]:
#         """Web arayÃ¼zÃ¼ iÃ§in geÃ§miÅŸ getir"""
#         target_agent = "ATLAS" if agent_name == "GENEL" else agent_name

#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute(
#                         """
#                         SELECT agent, role, message, timestamp
#                         FROM chat_history
#                         WHERE agent = ?
#                         ORDER BY id DESC
#                         LIMIT ?
#                         """,
#                         (target_agent, limit)
#                     )
#                     rows = cursor.fetchall()
#                     return [
#                         {
#                             "sender": "Siz" if row["role"] == "user" else row["agent"],
#                             "text": row["message"],
#                             "type": "user" if row["role"] == "user" else "agent",
#                             "time": row["timestamp"][11:16] if row["timestamp"] else ""
#                         }
#                         for row in reversed(rows)
#                     ]
#             except Exception as e:
#                 logger.error(f"Web geÃ§miÅŸi hatasÄ±: {e}")
#                 return []

#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
#     # CLEANUP & MANAGEMENT
#     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

#     def clear_history(self, only_chat: bool = True) -> bool:
#         """
#         HafÄ±zayÄ± temizle

#         Args:
#             only_chat: Sadece sohbet geÃ§miÅŸi mi (True) yoksa tÃ¼m hafÄ±za mÄ± (False)

#         Returns:
#             BaÅŸarÄ±lÄ±ysa True
#         """
#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute("DELETE FROM chat_history")
#                     if not only_chat:
#                         cursor.execute("DELETE FROM processed_files")
#                     conn.commit()

#                 if CHROMA_AVAILABLE and self.chroma_client:
#                     try:
#                         self.chroma_client.delete_collection("chat_history")
#                         self.history_collection = self.chroma_client.get_or_create_collection(
#                             name="chat_history",
#                             metadata={"hnsw:space": "cosine"},
#                             embedding_function=self.embedding_fn
#                         )

#                         if not only_chat:
#                             self.chroma_client.delete_collection("documents")
#                             self.docs_collection = self.chroma_client.get_or_create_collection(
#                                 name="documents",
#                                 metadata={"hnsw:space": "cosine"},
#                                 embedding_function=self.embedding_fn
#                             )

#                     except Exception as e:
#                         logger.warning(f"ChromaDB temizleme uyarÄ±sÄ±: {e}")

#                 mode = "Sohbet" if only_chat else "Tam"
#                 logger.info(f"ğŸ§¹ HafÄ±za temizlendi ({mode})")
#                 return True

#             except Exception as e:
#                 logger.error(f"HafÄ±za temizleme hatasÄ±: {e}")
#                 return False

#     def get_stats(self) -> Dict[str, Any]:
#         """HafÄ±za istatistiklerini getir"""
#         stats = {
#             "total_messages": 0,
#             "total_documents": 0,
#             "processed_files": 0,
#             "db_size_mb": 0,
#             "vector_db_enabled": CHROMA_AVAILABLE,
#             "embedding_backend": (
#                 f"Ollama ({Config.LOCAL_VEK})"
#                 if Config.AI_PROVIDER == "ollama"
#                 else f"SentenceTransformer ({MemoryConfig.SENTENCE_TRANSFORMER_MODEL})"
#             )
#         }

#         with self.lock:
#             try:
#                 with self._get_db_connection() as conn:
#                     cursor = conn.cursor()
#                     cursor.execute("SELECT COUNT(*) FROM chat_history")
#                     stats["total_messages"] = cursor.fetchone()[0]
#                     cursor.execute("SELECT COUNT(*) FROM processed_files")
#                     stats["processed_files"] = cursor.fetchone()[0]

#                 if self.db_path.exists():
#                     stats["db_size_mb"] = round(
#                         self.db_path.stat().st_size / (1024 * 1024), 2
#                     )

#                 if CHROMA_AVAILABLE and self.docs_collection:
#                     stats["total_documents"] = self.docs_collection.count()

#             except Exception as e:
#                 logger.error(f"Ä°statistik hatasÄ±: {e}")

#         return stats

#     def shutdown(self) -> None:
#         """Memory manager'Ä± kapat"""
#         logger.info("Memory Manager kapatÄ±lÄ±yor...")
#         if self.chroma_client:
#             try:
#                 self.chroma_client = None
#             except Exception as e:
#                 logger.warning(f"ChromaDB kapatma uyarÄ±sÄ±: {e}")
#         logger.info("âœ… Memory Manager kapatÄ±ldÄ±")


# # """
# # LotusAI Memory Management System
# # SÃ¼rÃ¼m: 2.5.3
# # AÃ§Ä±klama: Hybrid hafÄ±za sistemi (SQLite + ChromaDB) ve RAG implementasyonu

# # Ã–zellikler:
# # - KÄ±sa sÃ¼reli hafÄ±za: SQLite (hÄ±zlÄ± eriÅŸim)
# # - Uzun sÃ¼reli hafÄ±za: ChromaDB (anlamsal arama)
# # - RAG: Belge indeksleme ve retrieval
# # - GPU hÄ±zlandÄ±rma
# # - Thread-safe operasyonlar
# # """

# # import sqlite3
# # import logging
# # import threading
# # import hashlib
# # import re
# # from pathlib import Path
# # from datetime import datetime
# # from typing import List, Dict, Tuple, Any, Optional
# # from dataclasses import dataclass
# # from contextlib import contextmanager
# # from enum import Enum

# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # # CONFIG
# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # from config import Config

# # logger = logging.getLogger("LotusAI.Memory")


# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # # EXTERNAL LIBRARIES
# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # # PDF
# # try:
# #     import PyPDF2
# #     PDF_AVAILABLE = True
# # except ImportError:
# #     PDF_AVAILABLE = False
# #     logger.warning("âš ï¸ PyPDF2 yÃ¼klÃ¼ deÄŸil, PDF desteÄŸi yok")

# # # DOCX
# # try:
# #     import docx
# #     DOCX_AVAILABLE = True
# # except ImportError:
# #     DOCX_AVAILABLE = False
# #     logger.warning("âš ï¸ python-docx yÃ¼klÃ¼ deÄŸil, DOCX desteÄŸi yok")

# # # ChromaDB + Torch
# # try:
# #     import chromadb
# #     from chromadb.config import Settings
# #     from chromadb.utils import embedding_functions
# #     import torch
# #     CHROMA_AVAILABLE = True
# # except ImportError:
# #     CHROMA_AVAILABLE = False
# #     logger.warning("âš ï¸ ChromaDB/torch yÃ¼klÃ¼ deÄŸil, vektÃ¶r arama devre dÄ±ÅŸÄ±")


# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # # SABITLER
# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # class MemoryConfig:
# #     """Memory system konfigÃ¼rasyonu"""
# #     # Database
# #     DB_TIMEOUT = 10  # saniye
# #     DB_VERSION = 1
    
# #     # Chunking
# #     CHUNK_SIZE = 800  # karakter
# #     CHUNK_OVERLAP = 150
# #     MIN_CHUNK_LENGTH = 10
    
# #     # Retrieval
# #     MAX_RECENT_MESSAGES = 10
# #     MAX_HISTORY_RESULTS = 3
# #     MAX_DOCUMENT_RESULTS = 4
# #     MIN_MESSAGE_LENGTH_FOR_VECTOR = 15
    
# #     # Ingestion
# #     SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.md', '.docx']
    
# #     # Embedding
# #     EMBEDDING_MODEL = "all-MiniLM-L6-v2"
# #     # Alternatif (daha kaliteli): "paraphrase-multilingual-MiniLM-L12-v2"


# # class MessageRole(Enum):
# #     """Mesaj rolleri"""
# #     USER = "user"
# #     ASSISTANT = "assistant"
# #     MODEL = "model"
# #     SYSTEM = "system"


# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # # VERÄ° YAPILARI
# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # @dataclass
# # class ChatMessage:
# #     """Chat mesaj yapÄ±sÄ±"""
# #     role: str
# #     content: str
# #     agent: Optional[str] = None
# #     timestamp: Optional[datetime] = None


# # @dataclass
# # class ProcessedFile:
# #     """Ä°ÅŸlenmiÅŸ dosya bilgisi"""
# #     filename: str
# #     file_hash: str
# #     file_size: int
# #     processed_date: datetime


# # @dataclass
# # class DocumentChunk:
# #     """Belge parÃ§asÄ±"""
# #     content: str
# #     source: str
# #     chunk_id: int
# #     file_hash: str
# #     file_type: str


# # @dataclass
# # class MemoryContext:
# #     """HafÄ±za baÄŸlamÄ±"""
# #     recent_messages: List[Dict[str, str]]
# #     long_term_history: str
# #     relevant_documents: str


# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # # MEMORY MANAGER
# # # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# # class MemoryManager:
# #     """
# #     LotusAI Hybrid Memory System
    
# #     Mimari:
# #     â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”
# #     â”‚         Memory Manager                      â”‚
# #     â”œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¬â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¤
# #     â”‚   SQLite        â”‚   ChromaDB (Vector DB)    â”‚
# #     â”‚   (Fast Access) â”‚   (Semantic Search)       â”‚
# #     â”œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¼â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¤
# #     â”‚ â€¢ Chat History  â”‚ â€¢ Chat Embeddings         â”‚
# #     â”‚ â€¢ File Tracking â”‚ â€¢ Document Embeddings     â”‚
# #     â”‚                 â”‚ â€¢ GPU Accelerated         â”‚
# #     â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”´â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜
    
# #     Features:
# #     - Thread-safe operations
# #     - GPU-accelerated embeddings
# #     - RAG (Retrieval-Augmented Generation)
# #     - Multi-format document support
# #     - Smart text chunking
# #     - Deduplication (hash-based)
# #     """
    
# #     def __init__(self):
# #         """Memory manager baÅŸlatÄ±cÄ±"""
# #         # Paths
# #         self.work_dir = Config.WORK_DIR
# #         self.db_path = self.work_dir / "lotus_system.db"
# #         self.docs_path = self.work_dir / "documents"
# #         self.vector_db_path = self.work_dir / "lotus_vector_db"
        
# #         # Thread safety
# #         self.lock = threading.RLock()  # Reentrant lock
        
# #         # Create directories
# #         self.docs_path.mkdir(parents=True, exist_ok=True)
# #         self.vector_db_path.mkdir(parents=True, exist_ok=True)
        
# #         # Initialize SQLite
# #         self._init_sqlite()
        
# #         # ChromaDB components
# #         self.chroma_client: Optional[chromadb.Client] = None
# #         self.history_collection = None
# #         self.docs_collection = None
# #         self.embedding_fn = None
        
# #         # Initialize ChromaDB
# #         if CHROMA_AVAILABLE:
# #             self._init_chromadb()
        
# #         logger.info("âœ… Memory Manager baÅŸlatÄ±ldÄ±")
    
# #     def _init_sqlite(self) -> None:
# #         """SQLite veritabanÄ±nÄ± baÅŸlat"""
# #         with self.lock:
# #             try:
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
                    
# #                     # Chat history table
# #                     cursor.execute("""
# #                         CREATE TABLE IF NOT EXISTS chat_history (
# #                             id INTEGER PRIMARY KEY AUTOINCREMENT,
# #                             agent TEXT NOT NULL,
# #                             role TEXT NOT NULL,
# #                             message TEXT NOT NULL,
# #                             timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
# #                         )
# #                     """)
                    
# #                     # Processed files table
# #                     cursor.execute("""
# #                         CREATE TABLE IF NOT EXISTS processed_files (
# #                             filename TEXT PRIMARY KEY,
# #                             processed_date DATETIME NOT NULL,
# #                             file_hash TEXT NOT NULL,
# #                             file_size INTEGER NOT NULL
# #                         )
# #                     """)
                    
# #                     # Schema version table
# #                     cursor.execute("""
# #                         CREATE TABLE IF NOT EXISTS schema_version (
# #                             version INTEGER PRIMARY KEY,
# #                             applied_date DATETIME DEFAULT CURRENT_TIMESTAMP
# #                         )
# #                     """)
                    
# #                     # Insert schema version if not exists
# #                     cursor.execute(
# #                         "INSERT OR IGNORE INTO schema_version (version) VALUES (?)",
# #                         (MemoryConfig.DB_VERSION,)
# #                     )
                    
# #                     # Indexes for performance
# #                     cursor.execute(
# #                         "CREATE INDEX IF NOT EXISTS idx_chat_agent "
# #                         "ON chat_history(agent)"
# #                     )
# #                     cursor.execute(
# #                         "CREATE INDEX IF NOT EXISTS idx_chat_timestamp "
# #                         "ON chat_history(timestamp DESC)"
# #                     )
                    
# #                     conn.commit()
                
# #                 logger.info("âœ… SQLite veritabanÄ± hazÄ±r")
            
# #             except sqlite3.Error as e:
# #                 logger.error(f"âŒ SQLite baÅŸlatma hatasÄ±: {e}")
# #                 raise
    
# #     def _init_chromadb(self) -> None:
# #         """ChromaDB ve vektÃ¶r koleksiyonlarÄ±nÄ± baÅŸlat"""
# #         try:
# #             # Device selection
# #             if Config.USE_GPU and torch.cuda.is_available():
# #                 device = "cuda"
# #                 logger.info(f"ğŸš€ GPU aktif: {torch.cuda.get_device_name(0)}")
# #             else:
# #                 device = "cpu"
# #                 if Config.USE_GPU:
# #                     logger.warning("âš ï¸ GPU istendi ama CUDA yok, CPU kullanÄ±lÄ±yor")
            
# #             # Embedding function
# #             self.embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
# #                 model_name=MemoryConfig.EMBEDDING_MODEL,
# #                 device=device
# #             )
            
# #             # ChromaDB client
# #             self.chroma_client = chromadb.PersistentClient(
# #                 path=str(self.vector_db_path)
# #             )
            
# #             # Collections
# #             self.history_collection = self.chroma_client.get_or_create_collection(
# #                 name="chat_history",
# #                 metadata={"hnsw:space": "cosine"},
# #                 embedding_function=self.embedding_fn
# #             )
            
# #             self.docs_collection = self.chroma_client.get_or_create_collection(
# #                 name="documents",
# #                 metadata={"hnsw:space": "cosine"},
# #                 embedding_function=self.embedding_fn
# #             )
            
# #             doc_count = self.docs_collection.count()
# #             logger.info(f"âœ… ChromaDB aktif: {doc_count} belge parÃ§asÄ±")
        
# #         except Exception as e:
# #             logger.error(f"âŒ ChromaDB baÅŸlatma hatasÄ±: {e}")
# #             self.chroma_client = None
    
# #     @contextmanager
# #     def _get_db_connection(self):
# #         """
# #         Database connection context manager
        
# #         Yields:
# #             sqlite3.Connection
# #         """
# #         conn = None
# #         try:
# #             conn = sqlite3.connect(
# #                 str(self.db_path),
# #                 timeout=MemoryConfig.DB_TIMEOUT,
# #                 check_same_thread=False
# #             )
# #             conn.row_factory = sqlite3.Row
# #             yield conn
# #         finally:
# #             if conn:
# #                 conn.close()
    
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# #     # DOCUMENT INGESTION (RAG)
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
# #     def ingest_documents(self) -> str:
# #         """
# #         Belgeleri tara ve vektÃ¶r veritabanÄ±na indeksle
        
# #         Returns:
# #             Ä°ÅŸlem sonucu mesajÄ±
# #         """
# #         if not CHROMA_AVAILABLE or not self.docs_collection:
# #             return "âŒ VektÃ¶r veritabanÄ± aktif deÄŸil"
        
# #         new_files = 0
        
# #         for ext in MemoryConfig.SUPPORTED_EXTENSIONS:
# #             pattern = f"*{ext}"
            
# #             for file_path in self.docs_path.glob(pattern):
# #                 try:
# #                     if self._process_document(file_path):
# #                         new_files += 1
                
# #                 except Exception as e:
# #                     logger.error(f"Dosya iÅŸleme hatasÄ± ({file_path.name}): {e}")
        
# #         if new_files > 0:
# #             return f"âœ… {new_files} yeni belge indekslendi"
        
# #         return "â„¹ï¸ Yeni belge bulunamadÄ±"
    
# #     def _process_document(self, file_path: Path) -> bool:
# #         """
# #         Tek bir belgeyi iÅŸle
        
# #         Args:
# #             file_path: Dosya yolu
        
# #         Returns:
# #             Ä°ÅŸlem baÅŸarÄ±lÄ±ysa True
# #         """
# #         filename = file_path.name
        
# #         # File hash ve metadata
# #         file_content_raw = file_path.read_bytes()
# #         file_hash = hashlib.md5(file_content_raw).hexdigest()
# #         file_size = file_path.stat().st_size
        
# #         # Daha Ã¶nce iÅŸlendi mi?
# #         if self._is_file_processed(filename, file_hash):
# #             return False
        
# #         # Ä°Ã§eriÄŸi oku
# #         content = self._read_document_content(file_path)
        
# #         if not content or not content.strip():
# #             logger.warning(f"BoÅŸ iÃ§erik: {filename}")
# #             return False
        
# #         # Chunk'la
# #         chunks = self._smart_chunk_text(content)
        
# #         if not chunks:
# #             logger.warning(f"Chunk oluÅŸturulamadÄ±: {filename}")
# #             return False
        
# #         # VektÃ¶r DB'ye ekle
# #         ids = [
# #             f"{filename}_{i}_{file_hash[:8]}"
# #             for i in range(len(chunks))
# #         ]
        
# #         metadatas = [
# #             {
# #                 "source": filename,
# #                 "chunk_id": i,
# #                 "date": datetime.now().isoformat(),
# #                 "type": file_path.suffix[1:],
# #                 "hash": file_hash
# #             }
# #             for i in range(len(chunks))
# #         ]
        
# #         self.docs_collection.add(
# #             documents=chunks,
# #             metadatas=metadatas,
# #             ids=ids
# #         )
        
# #         # Ä°ÅŸlenmiÅŸ olarak iÅŸaretle
# #         self._mark_file_as_processed(filename, file_hash, file_size)
        
# #         logger.info(f"ğŸ“„ Ä°ndekslendi: {filename} ({len(chunks)} chunk)")
# #         return True
    
# #     def _read_document_content(self, file_path: Path) -> str:
# #         """
# #         Dosya tipine gÃ¶re iÃ§eriÄŸi oku
        
# #         Args:
# #             file_path: Dosya yolu
        
# #         Returns:
# #             Dosya iÃ§eriÄŸi
# #         """
# #         ext = file_path.suffix.lower()
        
# #         # Text/Markdown
# #         if ext in ['.txt', '.md']:
# #             return file_path.read_text(encoding='utf-8', errors='ignore')
        
# #         # PDF
# #         if ext == '.pdf' and PDF_AVAILABLE:
# #             content = []
# #             with open(file_path, 'rb') as f:
# #                 reader = PyPDF2.PdfReader(f)
# #                 for page in reader.pages:
# #                     text = page.extract_text()
# #                     if text:
# #                         content.append(text)
# #             return "\n".join(content)
        
# #         # DOCX
# #         if ext == '.docx' and DOCX_AVAILABLE:
# #             doc = docx.Document(file_path)
# #             return "\n".join([para.text for para in doc.paragraphs])
        
# #         return ""
    
# #     def _smart_chunk_text(
# #         self,
# #         text: str,
# #         chunk_size: Optional[int] = None,
# #         overlap: Optional[int] = None
# #     ) -> List[str]:
# #         """
# #         Metni anlamsal bÃ¼tÃ¼nlÃ¼ÄŸÃ¼ koruyarak chunk'la
        
# #         Args:
# #             text: Ä°ÅŸlenecek metin
# #             chunk_size: Chunk boyutu (karakter)
# #             overlap: Ã–rtÃ¼ÅŸme miktarÄ±
        
# #         Returns:
# #             Chunk listesi
# #         """
# #         if not text:
# #             return []
        
# #         chunk_size = chunk_size or MemoryConfig.CHUNK_SIZE
# #         overlap = overlap or MemoryConfig.CHUNK_OVERLAP
        
# #         # Normalize whitespace
# #         text = re.sub(r'\s+', ' ', text).strip()
        
# #         chunks = []
# #         start = 0
# #         text_len = len(text)
        
# #         while start < text_len:
# #             end = start + chunk_size
            
# #             # CÃ¼mle sonunda bitir (mÃ¼mkÃ¼nse)
# #             if end < text_len:
# #                 # Ã–nce chunk'Ä±n ortasÄ±ndan sonra ara
# #                 search_start = start + (chunk_size // 2)
                
# #                 best_break = -1
# #                 for punct in ['.', '!', '?', '\n']:
# #                     pos = text.rfind(punct, search_start, end)
# #                     if pos > best_break:
# #                         best_break = pos
                
# #                 if best_break != -1:
# #                     end = best_break + 1
            
# #             chunk = text[start:end].strip()
            
# #             if len(chunk) >= MemoryConfig.MIN_CHUNK_LENGTH:
# #                 chunks.append(chunk)
            
# #             # Overlap ile devam et
# #             start = max(start + 1, end - overlap)
            
# #             if end >= text_len:
# #                 break
        
# #         return chunks
    
# #     def _is_file_processed(self, filename: str, file_hash: str) -> bool:
# #         """
# #         DosyanÄ±n iÅŸlenip iÅŸlenmediÄŸini kontrol et
        
# #         Args:
# #             filename: Dosya adÄ±
# #             file_hash: Dosya hash'i
        
# #         Returns:
# #             Ä°ÅŸlendiyse True
# #         """
# #         with self.lock:
# #             try:
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
# #                     cursor.execute(
# #                         "SELECT 1 FROM processed_files "
# #                         "WHERE filename = ? AND file_hash = ?",
# #                         (filename, file_hash)
# #                     )
# #                     return cursor.fetchone() is not None
            
# #             except Exception as e:
# #                 logger.error(f"Dosya kontrolÃ¼ hatasÄ±: {e}")
# #                 return False
    
# #     def _mark_file_as_processed(
# #         self,
# #         filename: str,
# #         file_hash: str,
# #         file_size: int
# #     ) -> None:
# #         """
# #         DosyayÄ± iÅŸlenmiÅŸ olarak iÅŸaretle
        
# #         Args:
# #             filename: Dosya adÄ±
# #             file_hash: Dosya hash'i
# #             file_size: Dosya boyutu
# #         """
# #         with self.lock:
# #             try:
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
# #                     cursor.execute(
# #                         """
# #                         INSERT OR REPLACE INTO processed_files
# #                         (filename, processed_date, file_hash, file_size)
# #                         VALUES (?, ?, ?, ?)
# #                         """,
# #                         (filename, datetime.now().isoformat(), file_hash, file_size)
# #                     )
# #                     conn.commit()
            
# #             except Exception as e:
# #                 logger.error(f"Dosya iÅŸaretleme hatasÄ±: {e}")
    
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# #     # CHAT MEMORY
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
# #     def save(
# #         self,
# #         agent: str,
# #         role: str,
# #         message: str
# #     ) -> None:
# #         """
# #         MesajÄ± hafÄ±zaya kaydet
        
# #         Args:
# #             agent: Agent adÄ±
# #             role: Mesaj rolÃ¼ (user/assistant)
# #             message: Mesaj iÃ§eriÄŸi
# #         """
# #         if not message or not message.strip():
# #             return
        
# #         timestamp = datetime.now()
        
# #         # 1. SQLite'a kaydet
# #         with self.lock:
# #             try:
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
# #                     cursor.execute(
# #                         """
# #                         INSERT INTO chat_history (agent, role, message, timestamp)
# #                         VALUES (?, ?, ?, ?)
# #                         """,
# #                         (agent, role, message, timestamp.isoformat())
# #                     )
# #                     conn.commit()
            
# #             except sqlite3.Error as e:
# #                 logger.error(f"Mesaj kaydetme hatasÄ±: {e}")
        
# #         # 2. ChromaDB'ye vektÃ¶r olarak kaydet
# #         if (CHROMA_AVAILABLE and
# #             self.history_collection and
# #             len(message) >= MemoryConfig.MIN_MESSAGE_LENGTH_FOR_VECTOR):
            
# #             try:
# #                 unique_id = (
# #                     f"{agent}_{timestamp.timestamp()}_"
# #                     f"{hashlib.md5(message.encode()).hexdigest()[:6]}"
# #                 )
                
# #                 metadata = {
# #                     "agent": agent,
# #                     "role": role,
# #                     "time": timestamp.isoformat()
# #                 }
                
# #                 self.history_collection.add(
# #                     documents=[message],
# #                     metadatas=[metadata],
# #                     ids=[unique_id]
# #                 )
            
# #             except Exception as e:
# #                 logger.debug(f"VektÃ¶r kayÄ±t hatasÄ±: {e}")
    
# #     def load_context(
# #         self,
# #         agent: str,
# #         query: str,
# #         max_items: Optional[int] = None
# #     ) -> Tuple[List[Dict[str, str]], str, str]:
# #         """
# #         Agent iÃ§in baÄŸlam yÃ¼kle
        
# #         Args:
# #             agent: Agent adÄ±
# #             query: Sorgu metni
# #             max_items: Maksimum mesaj sayÄ±sÄ±
        
# #         Returns:
# #             Tuple[son mesajlar, uzun dÃ¶nem hafÄ±za, ilgili dokÃ¼manlar]
# #         """
# #         max_items = max_items or MemoryConfig.MAX_RECENT_MESSAGES
        
# #         # 1. Son mesajlar (SQLite)
# #         recent_messages = self._get_recent_messages(agent, max_items)
        
# #         # 2. Uzun dÃ¶nem hafÄ±za (ChromaDB - semantic search)
# #         long_term_history = self._search_history(agent, query)
        
# #         # 3. Ä°lgili dokÃ¼manlar (RAG)
# #         relevant_docs = self._search_documents(query)
        
# #         return recent_messages, long_term_history, relevant_docs
    
# #     def _get_recent_messages(
# #         self,
# #         agent: str,
# #         limit: int
# #     ) -> List[Dict[str, str]]:
# #         """
# #         Son mesajlarÄ± getir
        
# #         Args:
# #             agent: Agent adÄ±
# #             limit: Maksimum sayÄ±
        
# #         Returns:
# #             Mesaj listesi
# #         """
# #         with self.lock:
# #             try:
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
# #                     cursor.execute(
# #                         """
# #                         SELECT role, message
# #                         FROM chat_history
# #                         WHERE agent = ? OR agent = 'SYSTEM'
# #                         ORDER BY id DESC
# #                         LIMIT ?
# #                         """,
# #                         (agent, limit)
# #                     )
# #                     rows = cursor.fetchall()
                    
# #                     # Reverse to get chronological order
# #                     return [
# #                         {"role": row["role"], "content": row["message"]}
# #                         for row in reversed(rows)
# #                     ]
            
# #             except Exception as e:
# #                 logger.error(f"Son mesajlar getirme hatasÄ±: {e}")
# #                 return []
    
# #     def _search_history(self, agent: str, query: str) -> str:
# #         """
# #         GeÃ§miÅŸ sohbetlerde anlamsal arama
        
# #         Args:
# #             agent: Agent adÄ±
# #             query: Arama sorgusu
        
# #         Returns:
# #             Ä°lgili geÃ§miÅŸ mesajlar
# #         """
# #         if not CHROMA_AVAILABLE or not self.history_collection:
# #             return ""
        
# #         try:
# #             results = self.history_collection.query(
# #                 query_texts=[query],
# #                 n_results=MemoryConfig.MAX_HISTORY_RESULTS,
# #                 where={"agent": agent}
# #             )
            
# #             if results['documents'] and results['documents'][0]:
# #                 return "\n".join([
# #                     f"â€¢ {doc}"
# #                     for doc in results['documents'][0]
# #                 ])
        
# #         except Exception as e:
# #             logger.error(f"GeÃ§miÅŸ arama hatasÄ±: {e}")
        
# #         return ""
    
# #     def _search_documents(self, query: str) -> str:
# #         """
# #         Belgelerde anlamsal arama (RAG)
        
# #         Args:
# #             query: Arama sorgusu
        
# #         Returns:
# #             Ä°lgili belge parÃ§alarÄ±
# #         """
# #         if not CHROMA_AVAILABLE or not self.docs_collection:
# #             return ""
        
# #         try:
# #             results = self.docs_collection.query(
# #                 query_texts=[query],
# #                 n_results=MemoryConfig.MAX_DOCUMENT_RESULTS
# #             )
            
# #             if results['documents'] and results['documents'][0]:
# #                 docs_content = []
                
# #                 for i, doc_text in enumerate(results['documents'][0]):
# #                     source = results['metadatas'][0][i].get('source', 'Bilinmiyor')
# #                     docs_content.append(f"[{source}]: {doc_text}")
                
# #                 return "\n\n".join(docs_content)
        
# #         except Exception as e:
# #             logger.error(f"Belge arama hatasÄ±: {e}")
        
# #         return ""
    
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# #     # WEB INTERFACE
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
# #     def get_agent_history_for_web(
# #         self,
# #         agent_name: str,
# #         limit: int = 40
# #     ) -> List[Dict[str, Any]]:
# #         """
# #         Web arayÃ¼zÃ¼ iÃ§in geÃ§miÅŸ getir
        
# #         Args:
# #             agent_name: Agent adÄ±
# #             limit: Maksimum mesaj sayÄ±sÄ±
        
# #         Returns:
# #             FormatlanmÄ±ÅŸ mesaj listesi
# #         """
# #         target_agent = "ATLAS" if agent_name == "GENEL" else agent_name
        
# #         with self.lock:
# #             try:
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
# #                     cursor.execute(
# #                         """
# #                         SELECT agent, role, message, timestamp
# #                         FROM chat_history
# #                         WHERE agent = ?
# #                         ORDER BY id DESC
# #                         LIMIT ?
# #                         """,
# #                         (target_agent, limit)
# #                     )
# #                     rows = cursor.fetchall()
                    
# #                     return [
# #                         {
# #                             "sender": "Siz" if row["role"] == "user" else row["agent"],
# #                             "text": row["message"],
# #                             "type": "user" if row["role"] == "user" else "agent",
# #                             "time": row["timestamp"][11:16] if row["timestamp"] else ""
# #                         }
# #                         for row in reversed(rows)
# #                     ]
            
# #             except Exception as e:
# #                 logger.error(f"Web geÃ§miÅŸi hatasÄ±: {e}")
# #                 return []
    
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# #     # CLEANUP & MANAGEMENT
# #     # â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    
# #     def clear_history(self, only_chat: bool = True) -> bool:
# #         """
# #         HafÄ±zayÄ± temizle
        
# #         Args:
# #             only_chat: Sadece sohbet geÃ§miÅŸi mi (True) yoksa tÃ¼m hafÄ±za mÄ± (False)
        
# #         Returns:
# #             BaÅŸarÄ±lÄ±ysa True
# #         """
# #         with self.lock:
# #             try:
# #                 # SQLite temizliÄŸi
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
# #                     cursor.execute("DELETE FROM chat_history")
                    
# #                     if not only_chat:
# #                         cursor.execute("DELETE FROM processed_files")
                    
# #                     conn.commit()
                
# #                 # ChromaDB temizliÄŸi
# #                 if CHROMA_AVAILABLE and self.chroma_client:
# #                     try:
# #                         # Chat history collection
# #                         self.chroma_client.delete_collection("chat_history")
# #                         self.history_collection = self.chroma_client.get_or_create_collection(
# #                             name="chat_history",
# #                             metadata={"hnsw:space": "cosine"},
# #                             embedding_function=self.embedding_fn
# #                         )
                        
# #                         # Documents collection
# #                         if not only_chat:
# #                             self.chroma_client.delete_collection("documents")
# #                             self.docs_collection = self.chroma_client.get_or_create_collection(
# #                                 name="documents",
# #                                 metadata={"hnsw:space": "cosine"},
# #                                 embedding_function=self.embedding_fn
# #                             )
                    
# #                     except Exception as e:
# #                         logger.warning(f"ChromaDB temizleme uyarÄ±sÄ±: {e}")
                
# #                 mode = "Sohbet" if only_chat else "Tam"
# #                 logger.info(f"ğŸ§¹ HafÄ±za temizlendi ({mode})")
# #                 return True
            
# #             except Exception as e:
# #                 logger.error(f"HafÄ±za temizleme hatasÄ±: {e}")
# #                 return False
    
# #     def get_stats(self) -> Dict[str, Any]:
# #         """
# #         HafÄ±za istatistiklerini getir
        
# #         Returns:
# #             Ä°statistik dictionary'si
# #         """
# #         stats = {
# #             "total_messages": 0,
# #             "total_documents": 0,
# #             "processed_files": 0,
# #             "db_size_mb": 0,
# #             "vector_db_enabled": CHROMA_AVAILABLE
# #         }
        
# #         with self.lock:
# #             try:
# #                 # SQLite stats
# #                 with self._get_db_connection() as conn:
# #                     cursor = conn.cursor()
                    
# #                     # Message count
# #                     cursor.execute("SELECT COUNT(*) FROM chat_history")
# #                     stats["total_messages"] = cursor.fetchone()[0]
                    
# #                     # Processed files
# #                     cursor.execute("SELECT COUNT(*) FROM processed_files")
# #                     stats["processed_files"] = cursor.fetchone()[0]
                
# #                 # DB file size
# #                 if self.db_path.exists():
# #                     stats["db_size_mb"] = round(
# #                         self.db_path.stat().st_size / (1024 * 1024),
# #                         2
# #                     )
                
# #                 # ChromaDB stats
# #                 if CHROMA_AVAILABLE and self.docs_collection:
# #                     stats["total_documents"] = self.docs_collection.count()
            
# #             except Exception as e:
# #                 logger.error(f"Ä°statistik hatasÄ±: {e}")
        
# #         return stats
    
# #     def shutdown(self) -> None:
# #         """Memory manager'Ä± kapat"""
# #         logger.info("Memory Manager kapatÄ±lÄ±yor...")
        
# #         # ChromaDB'yi kapat (gerekirse)
# #         if self.chroma_client:
# #             try:
# #                 # ChromaDB otomatik olarak persist ediyor
# #                 self.chroma_client = None
# #             except Exception as e:
# #                 logger.warning(f"ChromaDB kapatma uyarÄ±sÄ±: {e}")
        
# #         logger.info("âœ… Memory Manager kapatÄ±ldÄ±")